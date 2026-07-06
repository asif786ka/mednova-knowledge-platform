"""
Document loaders for .md, .txt, .pdf, .docx.

PDF and DOCX loaders use optional libraries (pypdf, python-docx); if they are missing the
loader raises a clear, catchable error and the pipeline skips the file with a warning. The
core .md/.txt path has zero third-party dependencies.
"""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List

SUPPORTED = {".md", ".txt", ".pdf", ".docx"}


@dataclass
class LoadedDocument:
    document_id: str
    filename: str
    doc_type: str
    text: str


def _doc_type(name: str) -> str:
    n = name.lower()
    if "risk" in n:
        return "risk_assessment"
    if "meeting" in n:
        return "meeting_summary"
    if "voice" in n:
        return "voice_requirements"
    if "cloud" in n or "deployment" in n:
        return "cloud_deployment_plan"
    if "architecture" in n:
        return "architecture_notes"
    if "brief" in n:
        return "project_brief"
    if "service" in n or "insights" in n:
        return "service_description"
    if "design" in n or "platform" in n:
        return "design_document"
    return "document"


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"pypdf not installed, cannot read {path.name}: {exc}")
    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _read_docx(path: Path) -> str:
    try:
        import docx  # python-docx
    except Exception as exc:  # pragma: no cover
        raise RuntimeError(f"python-docx not installed, cannot read {path.name}: {exc}")
    document = docx.Document(str(path))
    return "\n".join(p.text for p in document.paragraphs)


def load_file(path: Path) -> LoadedDocument:
    path = Path(path)
    ext = path.suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"Unsupported file type: {ext}")
    if ext in {".md", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif ext == ".pdf":
        text = _read_pdf(path)
    else:  # .docx
        text = _read_docx(path)
    return LoadedDocument(
        document_id=path.stem,
        filename=path.name,
        doc_type=_doc_type(path.name),
        text=text,
    )


def load_folder(folder: Path) -> List[LoadedDocument]:
    folder = Path(folder)
    docs: List[LoadedDocument] = []
    for path in sorted(folder.rglob("*")):
        if path.is_file() and path.suffix.lower() in SUPPORTED:
            docs.append(load_file(path))
    return docs
